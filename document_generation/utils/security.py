import hashlib
import hmac
import secrets
import jwt
from datetime import datetime, timedelta
from typing import Optional, Dict, Any
from passlib.context import CryptContext
from cryptography.fernet import Fernet
import pyotp
import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor

from config.settings import settings
from utils.logger import get_logger

logger = get_logger(__name__)

# Password hashing
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")

# Encryption for sensitive data
encryption_key = settings.ENCRYPTION_KEY.encode() if settings.ENCRYPTION_KEY else Fernet.generate_key()
cipher_suite = Fernet(encryption_key)

def hash_password(password: str) -> str:
    """
    Hash a password using bcrypt
    """
    return pwd_context.hash(password)

def verify_password(plain_password: str, hashed_password: str) -> bool:
    """
    Verify a password against its hash
    """
    return pwd_context.verify(plain_password, hashed_password)

def create_access_token(data: Dict[str, Any], expires_delta: Optional[timedelta] = None) -> str:
    """
    Create a JWT access token
    """
    to_encode = data.copy()
    if expires_delta:
        expire = datetime.utcnow() + expires_delta
    else:
        expire = datetime.utcnow() + timedelta(minutes=settings.ACCESS_TOKEN_EXPIRE_MINUTES)
    
    to_encode.update({"exp": expire})
    encoded_jwt = jwt.encode(to_encode, settings.SECRET_KEY, algorithm=settings.ALGORITHM)
    return encoded_jwt

def verify_token(token: str) -> Optional[Dict[str, Any]]:
    """
    Verify and decode a JWT token
    """
    try:
        payload = jwt.decode(token, settings.SECRET_KEY, algorithms=[settings.ALGORITHM])
        return payload
    except jwt.PyJWTError as e:
        logger.warning(f"Token verification failed: {str(e)}")
        return None

def generate_api_key() -> str:
    """
    Generate a secure API key
    """
    return secrets.token_urlsafe(32)

def generate_totp_secret() -> str:
    """
    Generate a TOTP secret for 2FA
    """
    return pyotp.random_base32()

def verify_totp(secret: str, token: str) -> bool:
    """
    Verify a TOTP token
    """
    try:
        totp = pyotp.TOTP(secret)
        return totp.verify(token, valid_window=1)
    except Exception as e:
        logger.warning(f"TOTP verification failed: {str(e)}")
        return False

def encrypt_sensitive_data(data: str) -> str:
    """
    Encrypt sensitive data
    """
    try:
        encrypted_data = cipher_suite.encrypt(data.encode())
        return encrypted_data.decode()
    except Exception as e:
        logger.error(f"Encryption failed: {str(e)}")
        raise

def decrypt_sensitive_data(encrypted_data: str) -> str:
    """
    Decrypt sensitive data
    """
    try:
        decrypted_data = cipher_suite.decrypt(encrypted_data.encode())
        return decrypted_data.decode()
    except Exception as e:
        logger.error(f"Decryption failed: {str(e)}")
        raise

def generate_document_hash(file_path: str) -> str:
    """
    Generate SHA-256 hash of a document for integrity verification
    """
    try:
        hash_sha256 = hashlib.sha256()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_sha256.update(chunk)
        return hash_sha256.hexdigest()
    except Exception as e:
        logger.error(f"Document hashing failed: {str(e)}")
        raise

def verify_document_integrity(file_path: str, expected_hash: str) -> bool:
    """
    Verify document integrity using hash comparison
    """
    try:
        actual_hash = generate_document_hash(file_path)
        return hmac.compare_digest(actual_hash, expected_hash)
    except Exception as e:
        logger.error(f"Document integrity verification failed: {str(e)}")
        return False

def sanitize_filename(filename: str) -> str:
    """
    Sanitize filename to prevent path traversal attacks
    """
    # Remove path separators and dangerous characters
    dangerous_chars = ['/', '\\', '..', '<', '>', ':', '"', '|', '?', '*']
    sanitized = filename
    
    for char in dangerous_chars:
        sanitized = sanitized.replace(char, '_')
    
    # Limit length
    if len(sanitized) > 255:
        name, ext = os.path.splitext(sanitized)
        sanitized = name[:250] + ext
    
    return sanitized

def validate_file_type(file_path: str, allowed_extensions: list) -> bool:
    """
    Validate file type based on extension
    """
    file_extension = os.path.splitext(file_path)[1].lower()
    return file_extension in [ext.lower() for ext in allowed_extensions]

def add_watermark(file_path: str, watermark_text: str = "CONFIDENTIAL") -> str:
    """
    Add watermark to Word document
    """
    try:
        # Load document
        doc = Document(file_path)
        
        # Add watermark to each section
        for section in doc.sections:
            header = section.header
            watermark_paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            
            # Clear existing content
            watermark_paragraph.clear()
            
            # Add watermark text
            run = watermark_paragraph.add_run(watermark_text)
            run.font.color.rgb = RGBColor(192, 192, 192)  # Light gray
            run.font.size = Inches(0.5)
            watermark_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Save with watermark
        watermarked_path = file_path.replace('.docx', '_watermarked.docx')
        doc.save(watermarked_path)
        
        # Remove original file
        os.remove(file_path)
        
        return watermarked_path
        
    except Exception as e:
        logger.error(f"Watermark addition failed: {str(e)}")
        return file_path

def add_digital_signature(file_path: str) -> str:
    """
    Add digital signature to document (placeholder implementation)
    Note: Full digital signature implementation requires additional libraries
    and certificate management
    """
    try:
        # This is a simplified implementation
        # In production, you would use proper digital signature libraries
        
        # Generate signature metadata
        signature_data = {
            "signed_by": "ROTC Document System",
            "signed_at": datetime.utcnow().isoformat(),
            "document_hash": generate_document_hash(file_path)
        }
        
        # Load document
        doc = Document(file_path)
        
        # Add signature information as a footer
        for section in doc.sections:
            footer = section.footer
            signature_paragraph = footer.add_paragraph()
            signature_paragraph.add_run(
                f"Digitally signed by {signature_data['signed_by']} on {signature_data['signed_at']}"
            )
        
        # Save signed document
        signed_path = file_path.replace('.docx', '_signed.docx')
        doc.save(signed_path)
        
        # Remove original file
        os.remove(file_path)
        
        return signed_path
        
    except Exception as e:
        logger.error(f"Digital signature addition failed: {str(e)}")
        return file_path

def generate_secure_token(length: int = 32) -> str:
    """
    Generate a cryptographically secure random token
    """
    return secrets.token_urlsafe(length)

def validate_ip_address(ip_address: str, allowed_ips: list) -> bool:
    """
    Validate IP address against whitelist
    """
    if not allowed_ips:
        return True  # No restrictions if list is empty
    
    return ip_address in allowed_ips

def log_security_event(event_type: str, details: Dict[str, Any], user_id: Optional[str] = None):
    """
    Log security-related events
    """
    security_log = {
        "timestamp": datetime.utcnow().isoformat(),
        "event_type": event_type,
        "user_id": user_id,
        "details": details
    }
    
    logger.info(f"Security Event: {security_log}")

def anonymize_data(data: Dict[str, Any], fields_to_anonymize: list) -> Dict[str, Any]:
    """
    Anonymize sensitive fields in data
    """
    anonymized_data = data.copy()
    
    for field in fields_to_anonymize:
        if field in anonymized_data:
            if isinstance(anonymized_data[field], str):
                # Replace with asterisks, keeping first and last character
                value = anonymized_data[field]
                if len(value) > 2:
                    anonymized_data[field] = value[0] + '*' * (len(value) - 2) + value[-1]
                else:
                    anonymized_data[field] = '*' * len(value)
            else:
                anonymized_data[field] = "[REDACTED]"
    
    return anonymized_data

def validate_input_data(data: Dict[str, Any], schema: Dict[str, Any]) -> tuple[bool, list]:
    """
    Validate input data against schema to prevent injection attacks
    """
    errors = []
    
    for field, rules in schema.items():
        if field in data:
            value = data[field]
            
            # Check required fields
            if rules.get('required', False) and not value:
                errors.append(f"Field '{field}' is required")
                continue
            
            # Check data type
            expected_type = rules.get('type')
            if expected_type and not isinstance(value, expected_type):
                errors.append(f"Field '{field}' must be of type {expected_type.__name__}")
                continue
            
            # Check string length
            if isinstance(value, str):
                max_length = rules.get('max_length')
                if max_length and len(value) > max_length:
                    errors.append(f"Field '{field}' exceeds maximum length of {max_length}")
                
                min_length = rules.get('min_length')
                if min_length and len(value) < min_length:
                    errors.append(f"Field '{field}' is below minimum length of {min_length}")
                
                # Check for dangerous patterns
                dangerous_patterns = ['<script', 'javascript:', 'data:', 'vbscript:']
                if any(pattern in value.lower() for pattern in dangerous_patterns):
                    errors.append(f"Field '{field}' contains potentially dangerous content")
    
    return len(errors) == 0, errors

class RateLimiter:
    """
    Simple in-memory rate limiter
    """
    
    def __init__(self):
        self.requests = {}
    
    def is_allowed(self, identifier: str, max_requests: int, window_seconds: int) -> bool:
        """
        Check if request is allowed based on rate limiting
        """
        now = datetime.utcnow()
        window_start = now - timedelta(seconds=window_seconds)
        
        # Clean old entries
        if identifier in self.requests:
            self.requests[identifier] = [
                req_time for req_time in self.requests[identifier]
                if req_time > window_start
            ]
        else:
            self.requests[identifier] = []
        
        # Check if under limit
        if len(self.requests[identifier]) < max_requests:
            self.requests[identifier].append(now)
            return True
        
        return False

# Global rate limiter instance
rate_limiter = RateLimiter()