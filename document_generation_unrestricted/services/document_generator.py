import os
import asyncio
from datetime import datetime, timedelta
from typing import Dict, Any, List, Optional
from sqlalchemy import text
from sqlalchemy.orm import Session
from docx import Document
from docx.shared import Inches
from jinja2 import Environment, FileSystemLoader
import uuid
import hashlib
from pathlib import Path

from config.database import SessionLocal
from config.settings import settings
from models.cadet import CadetProfile, CadetStatus
from models.document import DocumentGeneration, DocumentTemplate, GenerationStatus
from utils.logger import get_logger
from utils.security import generate_document_hash, add_watermark, add_digital_signature

logger = get_logger(__name__)

class DocumentGeneratorService:
    """
    Service for generating documents from templates and cadet data
    """
    
    def __init__(self):
        self.template_env = Environment(
            loader=FileSystemLoader(settings.TEMPLATE_DIR),
            autoescape=True
        )
    
    @staticmethod
    async def generate_document_async(
        job_id: str,
        filters: Dict[str, Any],
        security_options: Dict[str, Any]
    ):
        """
        Asynchronous document generation task
        """
        db = SessionLocal()
        try:
            service = DocumentGeneratorService()
            await service.generate_document(
                db, job_id, filters, security_options
            )
        except Exception as e:
            logger.error(f"Document generation failed for job {job_id}: {str(e)}")
            # Update job status to failed
            generation_job = db.query(DocumentGeneration).filter(
                DocumentGeneration.job_id == job_id
            ).first()
            if generation_job:
                generation_job.status = GenerationStatus.FAILED
                generation_job.error_message = str(e)
                generation_job.completed_at = datetime.utcnow()
                db.commit()
        finally:
            db.close()
    
    async def generate_document(
        self,
        db: Session,
        job_id: str,
        filters: Dict[str, Any],
        security_options: Dict[str, Any]
    ):
        """
        Generate document based on job parameters
        """
        try:
            # Get generation job
            generation_job = db.query(DocumentGeneration).filter(
                DocumentGeneration.job_id == job_id
            ).first()
            
            if not generation_job:
                raise Exception(f"Generation job {job_id} not found")
            
            # Update status to processing
            generation_job.status = GenerationStatus.PROCESSING
            generation_job.started_at = datetime.utcnow()
            db.commit()
            
            # Get template
            template = db.query(DocumentTemplate).filter(
                DocumentTemplate.id == generation_job.template_id
            ).first()
            
            if not template:
                raise Exception(f"Template {generation_job.template_id} not found")
            
            # Get cadets based on filters
            cadets = await self._get_filtered_cadets(db, filters)
            
            if not cadets:
                raise Exception("No cadets found matching the filters")
            
            # Update total records
            generation_job.total_records = len(cadets)
            db.commit()
            
            # Generate document based on type
            # Map enum values to generator functions
            if generation_job.document_type.value == "summary":
                output_path = await self._generate_aer_report(
                    template, cadets, generation_job, db
                )
            elif generation_job.document_type.value == "custom":
                output_path = await self._generate_asr_report(
                    template, cadets, generation_job, db
                )
            elif generation_job.document_type.value in ("roster", "profiles"):
                output_path = await self._generate_cadet_list(
                    template, cadets, generation_job, db
                )
            else:
                raise Exception(f"Unsupported document type: {generation_job.document_type}")
            
            # Apply security options
            if security_options.get("watermark", True):
                output_path = await self._add_watermark(output_path)
            
            if security_options.get("password_protected", False):
                output_path = await self._add_password_protection(output_path)
            
            if security_options.get("digital_signature", False):
                output_path = await self._add_digital_signature(output_path)
            
            # Generate file hash for integrity
            file_hash = generate_document_hash(output_path)
            
            # Update job completion
            generation_job.status = GenerationStatus.COMPLETED
            generation_job.output_path = output_path
            generation_job.output_filename = os.path.basename(output_path)
            generation_job.file_size = os.path.getsize(output_path)
            generation_job.file_hash = file_hash
            generation_job.completed_at = datetime.utcnow()
            generation_job.processing_time = int(
                (generation_job.completed_at - generation_job.started_at).total_seconds()
            )
            generation_job.access_expires_at = datetime.utcnow() + timedelta(
                days=settings.DOCUMENT_ACCESS_EXPIRY_DAYS
            )
            generation_job.progress_percentage = 100.0
            generation_job.processed_records = len(cadets)
            
            db.commit()
            
            logger.info(f"Document generation completed for job {job_id}")
            
        except Exception as e:
            logger.error(f"Document generation error for job {job_id}: {str(e)}")
            raise
    
    async def _get_filtered_cadets(
        self, db: Session, filters: Dict[str, Any]
    ) -> List[CadetProfile]:
        """
        Get cadets based on filters from the existing ROTC database using a raw SQL join.
        UNRESTRICTED VERSION: Bypasses all approval_status and user status checks.
        Includes ALL cadets regardless of approval status, user status, or role.
        Returns a list of lightweight objects with a to_dict() method compatible with templates.
        """
        conditions = [
            "cp.user_id = u.id"  # Only require the join, no status filtering
        ]
        params: Dict[str, Any] = {}

        if filters.get("year_level"):
            conditions.append("u.year_level = :year_level")
            params["year_level"] = filters["year_level"]
        if filters.get("course"):
            conditions.append("cp.course = :course")
            params["course"] = filters["course"]
        if filters.get("platoon"):
            conditions.append("cp.platoon = :platoon")
            params["platoon"] = filters["platoon"]
        if filters.get("semester"):
            conditions.append("cp.semester = :semester")
            params["semester"] = filters["semester"]
        if filters.get("academic_year"):
            conditions.append("cp.academic_year = :academic_year")
            params["academic_year"] = filters["academic_year"]

        where_clause = " AND ".join(conditions)
        sql = text(f"""
            SELECT 
                cp.id,
                u.id AS user_id,
                u.username,
                u.role,
                u.approval_status,
                u.status AS user_status,
                cp.first_name,
                cp.last_name,
                cp.middle_name,
                cp.contact_number,
                cp.email,
                cp.course,
                cp.year_level,
                cp.platoon,
                cp.status,
                cp.semester,
                cp.academic_year,
                cp.address,
                cp.emergency_contact,
                cp.emergency_phone,
                cp.created_at,
                cp.updated_at
            FROM users u
            JOIN cadet_profiles cp ON cp.user_id = u.id
            WHERE {where_clause}
            ORDER BY cp.last_name, cp.first_name
        """)

        result = db.execute(sql, params)

        class RowAdapter:
            def __init__(self, row: Dict[str, Any]):
                self.row = row
            def to_dict(self):
                # Match CadetProfile.to_dict keys used in templates
                # UNRESTRICTED VERSION: Include approval_status and user_status fields
                return {
                    "id": self.row.get("id"),
                    "student_id": self.row.get("user_id"),
                    "first_name": self.row.get("first_name"),
                    "last_name": self.row.get("last_name"),
                    "middle_name": self.row.get("middle_name"),
                    "full_name": f"{self.row.get('first_name','')} {self.row.get('last_name','')}".strip(),
                    "contact_number": self.row.get("contact_number"),
                    "email": self.row.get("email"),
                    "course": self.row.get("course"),
                    "year_level": self.row.get("year_level"),
                    "platoon": self.row.get("platoon"),
                    "status": self.row.get("status"),
                    "user_role": self.row.get("role"),
                    "approval_status": self.row.get("approval_status"),
                    "user_status": self.row.get("user_status"),
                    "semester": self.row.get("semester"),
                    "academic_year": self.row.get("academic_year"),
                    "address": self.row.get("address"),
                    "emergency_contact": self.row.get("emergency_contact"),
                    "emergency_phone": self.row.get("emergency_phone"),
                    "created_at": self.row.get("created_at").isoformat() if self.row.get("created_at") else None,
                    "updated_at": self.row.get("updated_at").isoformat() if self.row.get("updated_at") else None,
                }

        # SQLAlchemy Row objects can be converted to dict via _mapping
        cadet_rows = [RowAdapter(dict(r._mapping)) for r in result]
        return cadet_rows
    
    async def _generate_aer_report(
        self,
        template: DocumentTemplate,
        cadets: List[CadetProfile],
        generation_job: DocumentGeneration,
        db: Session
    ) -> str:
        """
        Generate AER (Annual Enrollment Report)
        """
        # Load template
        doc = Document(template.file_path)
        
        # Prepare data for template
        template_data = {
            "report_title": "Annual Enrollment Report (AER)",
            "generation_date": datetime.now().strftime("%B %d, %Y"),
            "academic_year": generation_job.filters.get("academic_year", "Current"),
            "total_cadets": len(cadets),
            "cadets": [cadet.to_dict() for cadet in cadets],
            "summary": self._generate_aer_summary(cadets)
        }
        
        # Replace placeholders in document
        await self._replace_template_placeholders(doc, template_data)
        
        # Generate output filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"AER_Report_{timestamp}_{generation_job.job_id[:8]}.docx"
        output_path = os.path.join(settings.OUTPUT_DIR, filename)
        
        # Save document
        doc.save(output_path)
        
        return output_path
    
    async def _generate_asr_report(
        self,
        template: DocumentTemplate,
        cadets: List[CadetProfile],
        generation_job: DocumentGeneration,
        db: Session
    ) -> str:
        """
        Generate ASR (Annual Statistical Report)
        """
        # Load template
        doc = Document(template.file_path)
        
        # Prepare statistical data
        stats = self._calculate_asr_statistics(cadets)
        
        template_data = {
            "report_title": "Annual Statistical Report (ASR)",
            "generation_date": datetime.now().strftime("%B %d, %Y"),
            "academic_year": generation_job.filters.get("academic_year", "Current"),
            "statistics": stats,
            "total_cadets": len(cadets)
        }
        
        # Replace placeholders in document
        await self._replace_template_placeholders(doc, template_data)
        
        # Generate output filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"ASR_Report_{timestamp}_{generation_job.job_id[:8]}.docx"
        output_path = os.path.join(settings.OUTPUT_DIR, filename)
        
        # Save document
        doc.save(output_path)
        
        return output_path
    
    async def _generate_cadet_list(
        self,
        template: DocumentTemplate,
        cadets: List[CadetProfile],
        generation_job: DocumentGeneration,
        db: Session
    ) -> str:
        """
        Generate Cadet List document
        """
        # Load template
        doc = Document(template.file_path)
        
        # Sort cadets by last name
        sorted_cadets = sorted(cadets, key=lambda c: c.last_name)
        
        template_data = {
            "report_title": "ROTC Cadet List",
            "generation_date": datetime.now().strftime("%B %d, %Y"),
            "semester": generation_job.filters.get("semester", "Current"),
            "academic_year": generation_job.filters.get("academic_year", "Current"),
            "total_cadets": len(cadets),
            "cadets": [cadet.to_dict() for cadet in sorted_cadets]
        }
        
        # Replace placeholders in document
        await self._replace_template_placeholders(doc, template_data)
        
        # Generate output filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"Cadet_List_{timestamp}_{generation_job.job_id[:8]}.docx"
        output_path = os.path.join(settings.OUTPUT_DIR, filename)
        
        # Save document
        doc.save(output_path)
        
        return output_path
    
    async def _replace_template_placeholders(
        self, doc: Document, data: Dict[str, Any]
    ):
        """
        Replace placeholders in Word document with actual data
        """
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            for key, value in data.items():
                if f"{{{{{key}}}}}" in paragraph.text:
                    paragraph.text = paragraph.text.replace(
                        f"{{{{{key}}}}}", str(value)
                    )
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in data.items():
                        if f"{{{{{key}}}}}" in cell.text:
                            cell.text = cell.text.replace(
                                f"{{{{{key}}}}}", str(value)
                            )
    
    def _generate_aer_summary(self, cadets: List[CadetProfile]) -> Dict[str, Any]:
        """
        Generate summary statistics for AER report
        """
        total = len(cadets)
        
        # Count by year level
        year_counts = {}
        for cadet in cadets:
            year = cadet.year_level.value if cadet.year_level else "Unknown"
            year_counts[year] = year_counts.get(year, 0) + 1
        
        # Count by course
        course_counts = {}
        for cadet in cadets:
            course = cadet.course or "Unknown"
            course_counts[course] = course_counts.get(course, 0) + 1
        
        # Count by status
        status_counts = {}
        for cadet in cadets:
            status = cadet.status.value if cadet.status else "Unknown"
            status_counts[status] = status_counts.get(status, 0) + 1
        
        return {
            "total_cadets": total,
            "by_year_level": year_counts,
            "by_course": course_counts,
            "by_status": status_counts
        }
    
    def _calculate_asr_statistics(self, cadets: List[CadetProfile]) -> Dict[str, Any]:
        """
        Calculate detailed statistics for ASR report
        """
        total = len(cadets)
        
        # Gender distribution
        gender_counts = {"Male": 0, "Female": 0, "Other": 0}
        for cadet in cadets:
            gender = cadet.gender or "Other"
            if gender.lower() in ["m", "male"]:
                gender_counts["Male"] += 1
            elif gender.lower() in ["f", "female"]:
                gender_counts["Female"] += 1
            else:
                gender_counts["Other"] += 1
        
        # Age distribution
        age_groups = {"18-20": 0, "21-23": 0, "24+": 0}
        current_year = datetime.now().year
        
        for cadet in cadets:
            if cadet.birth_date:
                age = current_year - cadet.birth_date.year
                if 18 <= age <= 20:
                    age_groups["18-20"] += 1
                elif 21 <= age <= 23:
                    age_groups["21-23"] += 1
                else:
                    age_groups["24+"] += 1
        
        # Academic performance (if GPA available)
        gpa_stats = {"high": 0, "medium": 0, "low": 0}
        for cadet in cadets:
            if hasattr(cadet, 'gpa') and cadet.gpa:
                if cadet.gpa >= 3.5:
                    gpa_stats["high"] += 1
                elif cadet.gpa >= 2.5:
                    gpa_stats["medium"] += 1
                else:
                    gpa_stats["low"] += 1
        
        return {
            "total_cadets": total,
            "gender_distribution": gender_counts,
            "age_distribution": age_groups,
            "academic_performance": gpa_stats,
            "enrollment_trends": self._calculate_enrollment_trends(cadets)
        }
    
    def _calculate_enrollment_trends(self, cadets: List[CadetProfile]) -> Dict[str, Any]:
        """
        Calculate enrollment trends
        """
        # Group by enrollment date if available
        monthly_enrollment = {}
        for cadet in cadets:
            if cadet.created_at:
                month_key = cadet.created_at.strftime("%Y-%m")
                monthly_enrollment[month_key] = monthly_enrollment.get(month_key, 0) + 1
        
        return {
            "monthly_enrollment": monthly_enrollment,
            "total_new_enrollments": len(cadets)
        }
    
    async def _add_watermark(self, file_path: str) -> str:
        """
        Add watermark to document
        """
        try:
            return add_watermark(file_path, "ROTC CONFIDENTIAL")
        except Exception as e:
            logger.warning(f"Failed to add watermark: {str(e)}")
            return file_path
    
    async def _add_password_protection(self, file_path: str) -> str:
        """
        Add password protection to document
        """
        try:
            # Generate random password
            password = str(uuid.uuid4())[:12]
            # Note: Password protection for DOCX requires additional libraries
            # This is a placeholder for the actual implementation
            logger.info(f"Document password: {password}")
            return file_path
        except Exception as e:
            logger.warning(f"Failed to add password protection: {str(e)}")
            return file_path
    
    async def _add_digital_signature(self, file_path: str) -> str:
        """
        Add digital signature to document
        """
        try:
            return add_digital_signature(file_path)
        except Exception as e:
            logger.warning(f"Failed to add digital signature: {str(e)}")
            return file_path